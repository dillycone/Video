import sys
import json
import base64
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.style.font.color.rgb = RGBColor(204, 0, 0)  # Match the red color theme

def add_image(doc, image_data, width_inches=6.0):
    # Remove data URL prefix if present
    if image_data.startswith('data:image'):
        image_data = image_data.split(',')[1]
    
    # Decode base64 image
    image_bytes = base64.b64decode(image_data)
    image_stream = io.BytesIO(image_bytes)
    image = Image.open(image_stream)
    
    # Save to temporary stream in PNG format
    temp_stream = io.BytesIO()
    image.save(temp_stream, format='PNG')
    temp_stream.seek(0)
    
    # Add to document
    doc.add_picture(temp_stream, width=Inches(width_inches))

def generate_word_doc(data):
    procedure = data['procedure']
    image_scale = data['imageScale'] / 100.0  # Convert percentage to decimal
    
    doc = Document()
    
    # Title
    add_heading(doc, procedure['title'])
    doc.add_paragraph()
    
    # Overview
    add_heading(doc, 'Overview', 2)
    doc.add_paragraph(procedure['overview'])
    doc.add_paragraph()
    
    # Prerequisites
    add_heading(doc, 'Prerequisites', 2)
    for prereq in procedure['prerequisites']:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(prereq)
    doc.add_paragraph()
    
    # Procedure Steps
    add_heading(doc, 'Procedure', 2)
    for i, step in enumerate(procedure['steps'], 1):
        # Main step
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        p.add_run(step['main'])
        
        # Sub-steps
        for sub in step['sub']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run('• ').bold = True
            p.add_run(sub)
        
        # Warnings
        for warning in step['warnings']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run('⚠️ ').bold = True
            p.add_run(warning)
        
        # Tips
        for tip in step['tips']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run('💡 ').bold = True
            p.add_run(tip)
        
        # Images
        if step.get('frames'):
            for frame in step['frames']:
                add_image(doc, frame['image'], width_inches=6.0 * image_scale)
                caption = doc.add_paragraph(f"Time: {frame['timestamp']}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    # Verification
    add_heading(doc, 'Verification', 2)
    doc.add_paragraph(procedure['verification'])
    doc.add_paragraph()
    
    # Troubleshooting
    add_heading(doc, 'Troubleshooting', 2)
    for item in procedure['troubleshooting']:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(item)
    
    # Save to memory stream
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream.getvalue()

if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python generate_word.py <procedure_data.json>")
        sys.exit(1)
    
    with open(sys.argv[1], 'r') as f:
        data = json.load(f)
    
    doc_bytes = generate_word_doc(data)
    sys.stdout.buffer.write(doc_bytes) 